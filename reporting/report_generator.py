from docx import Document
from datetime import datetime
from pathlib import Path
from typing import List, Dict

class ReportGenerator:
    def generate(self, results: List[Dict], output_dir: str) -> str:
        doc = Document()
        doc.add_heading("UAT vs Production Comparison Report", level=1)
        doc.add_paragraph(f"Generated: {datetime.now()}")

        passed = sum(1 for r in results if r["status"] == "PASS")
        failed = sum(1 for r in results if r["status"] == "FAIL")
        doc.add_paragraph(f"Total pages: {len(results)}")
        doc.add_paragraph(f"Passed: {passed}")
        doc.add_paragraph(f"Failed: {failed}")

        # Per-page details
        for r in results:
            doc.add_heading(r["test_name"], level=2)
            doc.add_paragraph(f"Status: {r['status']}")
            doc.add_paragraph(f"UAT: {r['uat_url']}")
            doc.add_paragraph(f"Prod: {r['prod_url']}")

            # Content section
            doc.add_heading("Content comparison", level=3)
            content_issues = r.get("content_issues", [])
            if content_issues:
                doc.add_paragraph("UAT vs Prod text differences:")
                for d in content_issues:
                    side = "UAT only" if d["side"] == "UAT" else "Prod only"
                    doc.add_paragraph(f"{side}: {d['text']}", style="List Bullet")
            else:
                doc.add_paragraph("Content match: no text differences found between UAT and Prod.")

            # Links section
            doc.add_heading("Link comparison", level=3)
            link_issues = r.get("link_issues", [])
            if link_issues:
                for issue in link_issues:
                    doc.add_paragraph(issue, style="List Bullet")
            else:
                doc.add_paragraph("Links match: no link differences detected or link checks not implemented yet.")

            # Images section
            doc.add_heading("Image comparison", level=3)
            image_issues = r.get("image_issues", [])
            if image_issues:
                for issue in image_issues:
                    doc.add_paragraph(issue, style="List Bullet")
            else:
                doc.add_paragraph("Images match: no image differences detected or image checks not implemented yet.")

        Path(output_dir).mkdir(exist_ok=True, parents=True)
        path = Path(output_dir) / f"UAT_vs_Prod_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(str(path))
        return str(path)
